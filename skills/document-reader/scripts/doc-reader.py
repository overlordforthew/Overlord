#!/usr/bin/env python3
"""
Document Reader — Extract text, summarize, search, and convert documents.
Supports: PDF, XLSX, CSV, DOCX, TXT, JSON
"""

import argparse
import csv
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Auto-install missing pip packages on first run
# ---------------------------------------------------------------------------
try:
    import openpyxl
except ImportError:
    subprocess.check_call(
        ['pip3', 'install', '--break-system-packages', 'openpyxl', 'python-docx', 'PyPDF2'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    import openpyxl

try:
    import docx
except ImportError:
    subprocess.check_call(
        ['pip3', 'install', '--break-system-packages', 'python-docx'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    import docx

try:
    import PyPDF2
except ImportError:
    subprocess.check_call(
        ['pip3', 'install', '--break-system-packages', 'PyPDF2'],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    import PyPDF2

# ---------------------------------------------------------------------------
# File type detection
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {'.pdf', '.xlsx', '.csv', '.docx', '.txt', '.json'}


def detect_type(filepath: str) -> str:
    """Return the normalised extension (e.g. '.pdf')."""
    ext = Path(filepath).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    return ext


def _require_file(filepath: str) -> None:
    if not os.path.isfile(filepath):
        print(f"Error: file not found — {filepath}", file=sys.stderr)
        sys.exit(1)

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def extract_pdf(filepath: str) -> str:
    """Extract text from PDF. Prefer pdftotext; fall back to PyPDF2."""
    if shutil.which('pdftotext'):
        result = subprocess.run(
            ['pdftotext', '-layout', filepath, '-'],
            capture_output=True, text=True,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    # Fallback
    text_parts = []
    with open(filepath, 'rb') as fh:
        reader = PyPDF2.PdfReader(fh)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return '\n'.join(text_parts)


def extract_xlsx(filepath: str, max_rows: int = 20) -> str:
    """Extract text from XLSX — sheet names, first N rows per sheet, counts."""
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        total_rows = len(rows)
        total_cols = ws.max_column or 0
        lines.append(f"\n=== Sheet: {sheet_name} ({total_rows} rows, {total_cols} cols) ===\n")
        for i, row in enumerate(rows[:max_rows]):
            formatted = '\t'.join(str(c) if c is not None else '' for c in row)
            lines.append(formatted)
        if total_rows > max_rows:
            lines.append(f"  ... ({total_rows - max_rows} more rows)")
    wb.close()
    return '\n'.join(lines)


def extract_csv(filepath: str) -> str:
    """Extract text from CSV."""
    lines = []
    with open(filepath, newline='', encoding='utf-8', errors='replace') as fh:
        reader = csv.reader(fh)
        for row in reader:
            lines.append('\t'.join(row))
    return '\n'.join(lines)


def extract_docx(filepath: str) -> str:
    """Extract text from DOCX."""
    doc = docx.Document(filepath)
    return '\n'.join(p.text for p in doc.paragraphs)


def extract_txt(filepath: str) -> str:
    """Read plain text or JSON."""
    with open(filepath, encoding='utf-8', errors='replace') as fh:
        return fh.read()


EXTRACTORS = {
    '.pdf': extract_pdf,
    '.xlsx': extract_xlsx,
    '.csv': extract_csv,
    '.docx': extract_docx,
    '.txt': extract_txt,
    '.json': extract_txt,
}


def extract_text(filepath: str) -> str:
    ext = detect_type(filepath)
    return EXTRACTORS[ext](filepath)

# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------


def cmd_read(args):
    """Extract and print text from a document."""
    _require_file(args.file)
    text = extract_text(args.file)
    print(text)


def cmd_summary(args):
    """Extract text, truncate, and generate AI summary via llm CLI."""
    _require_file(args.file)
    text = extract_text(args.file)
    # Truncate to ~3000 chars for the LLM
    truncated = text[:3000]
    if len(text) > 3000:
        truncated += "\n\n[... truncated ...]"

    prompt = f"Summarize this document concisely:\n\n{truncated}"

    try:
        result = subprocess.run(
            ['llm', '-m', 'openrouter/openrouter/free', prompt],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode == 0:
            print(result.stdout.strip())
        else:
            print(f"LLM error: {result.stderr.strip()}", file=sys.stderr)
            sys.exit(1)
    except FileNotFoundError:
        print("Error: 'llm' CLI not found. Install with: pip install llm", file=sys.stderr)
        sys.exit(1)
    except subprocess.TimeoutExpired:
        print("Error: LLM summarization timed out after 120s", file=sys.stderr)
        sys.exit(1)


def cmd_search(args):
    """Search for a query string within a document."""
    _require_file(args.file)
    text = extract_text(args.file)
    query = args.query.lower()
    lines = text.splitlines()
    matches = []
    for i, line in enumerate(lines, 1):
        if query in line.lower():
            matches.append((i, line))

    if not matches:
        print(f"No matches found for '{args.query}'")
        return

    print(f"Found {len(matches)} match(es) for '{args.query}':\n")
    for line_num, line in matches:
        print(f"  Line {line_num}: {line.strip()}")


def cmd_info(args):
    """Show file metadata: type, size, pages/sheets/rows count."""
    _require_file(args.file)
    filepath = args.file
    ext = detect_type(filepath)
    stat = os.stat(filepath)
    size_bytes = stat.st_size

    # Human-readable size
    if size_bytes < 1024:
        size_str = f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        size_str = f"{size_bytes / 1024:.1f} KB"
    else:
        size_str = f"{size_bytes / (1024 * 1024):.1f} MB"

    print(f"File:  {os.path.basename(filepath)}")
    print(f"Path:  {os.path.abspath(filepath)}")
    print(f"Type:  {ext.lstrip('.')}")
    print(f"Size:  {size_str} ({size_bytes:,} bytes)")

    if ext == '.pdf':
        try:
            with open(filepath, 'rb') as fh:
                reader = PyPDF2.PdfReader(fh)
                print(f"Pages: {len(reader.pages)}")
        except Exception as e:
            print(f"Pages: (error reading — {e})")

    elif ext == '.xlsx':
        try:
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            print(f"Sheets: {len(wb.sheetnames)} — {', '.join(wb.sheetnames)}")
            for name in wb.sheetnames:
                ws = wb[name]
                rows = sum(1 for _ in ws.iter_rows(values_only=True))
                cols = ws.max_column or 0
                print(f"  {name}: {rows} rows, {cols} cols")
            wb.close()
        except Exception as e:
            print(f"Sheets: (error reading — {e})")

    elif ext == '.csv':
        try:
            with open(filepath, newline='', encoding='utf-8', errors='replace') as fh:
                row_count = sum(1 for _ in csv.reader(fh))
            print(f"Rows:  {row_count}")
        except Exception as e:
            print(f"Rows:  (error reading — {e})")

    elif ext == '.docx':
        try:
            doc = docx.Document(filepath)
            print(f"Paragraphs: {len(doc.paragraphs)}")
        except Exception as e:
            print(f"Paragraphs: (error reading — {e})")

    elif ext in ('.txt', '.json'):
        try:
            with open(filepath, encoding='utf-8', errors='replace') as fh:
                content = fh.read()
            line_count = content.count('\n') + (1 if content and not content.endswith('\n') else 0)
            print(f"Lines: {line_count}")
            if ext == '.json':
                try:
                    parsed = json.loads(content)
                    if isinstance(parsed, list):
                        print(f"JSON items: {len(parsed)}")
                    elif isinstance(parsed, dict):
                        print(f"JSON keys: {len(parsed)}")
                except json.JSONDecodeError:
                    print("JSON: invalid")
        except Exception as e:
            print(f"Lines: (error reading — {e})")


def cmd_convert(args):
    """Convert a document to plain text and write to file or stdout."""
    _require_file(args.file)
    if args.to != 'txt':
        print(f"Error: only '--to txt' is currently supported", file=sys.stderr)
        sys.exit(1)

    text = extract_text(args.file)
    out_path = Path(args.file).with_suffix('.txt')

    # Avoid overwriting the source if it's already .txt
    if Path(args.file).suffix.lower() == '.txt':
        print("File is already .txt — nothing to convert.")
        return

    with open(out_path, 'w', encoding='utf-8') as fh:
        fh.write(text)

    print(f"Converted to: {out_path}")

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description='Document Reader — read, summarize, search, and convert documents.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Supported formats: PDF, XLSX, CSV, DOCX, TXT, JSON\n\n"
            "Examples:\n"
            "  doc-reader.py read   report.pdf\n"
            "  doc-reader.py summary spreadsheet.xlsx\n"
            "  doc-reader.py search contract.docx 'payment terms'\n"
            "  doc-reader.py info   data.csv\n"
            "  doc-reader.py convert invoice.pdf --to txt\n"
        ),
    )
    subparsers = parser.add_subparsers(dest='command', required=True)

    # read
    p_read = subparsers.add_parser('read', help='Extract and print text from a file')
    p_read.add_argument('file', help='Path to the document')
    p_read.set_defaults(func=cmd_read)

    # summary
    p_summary = subparsers.add_parser('summary', help='Extract text and generate AI summary')
    p_summary.add_argument('file', help='Path to the document')
    p_summary.set_defaults(func=cmd_summary)

    # search
    p_search = subparsers.add_parser('search', help='Search for text within a document')
    p_search.add_argument('file', help='Path to the document')
    p_search.add_argument('query', help='Text to search for')
    p_search.set_defaults(func=cmd_search)

    # info
    p_info = subparsers.add_parser('info', help='Show file metadata')
    p_info.add_argument('file', help='Path to the document')
    p_info.set_defaults(func=cmd_info)

    # convert
    p_convert = subparsers.add_parser('convert', help='Convert document to plain text')
    p_convert.add_argument('file', help='Path to the document')
    p_convert.add_argument('--to', default='txt', choices=['txt'], help='Target format (default: txt)')
    p_convert.set_defaults(func=cmd_convert)

    args = parser.parse_args()
    args.func(args)


if __name__ == '__main__':
    main()
